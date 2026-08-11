"""文档转换核心逻辑（移植自 puremark-converter / main_server）。

路由策略：
    office 文档（docx/xlsx/pptx/txt/csv/md 等）→ 本地 MarkItDown 提取
    图片（png/jpg/bmp/tif/webp 等）           → OCR：优先转发配置的 OCR 节点，
                                                 节点不可达时降级为本地 PaddleOCR（可选安装）
    PDF                                        → 先分类：文本型走 MarkItDown，
                                                 扫描型走 OCR（同上含降级）

可选能力：
    crop            按比例裁剪区域后再转换（图片 / 文本型 PDF / 扫描型 PDF）
    output_format   md（默认）/ txt / docx 三种输出
    layout_reconstruction  基于 OCR 行级 bbox 的几何启发式版面重建（标题/段落/缩进）
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import shutil
import subprocess
import threading
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import fitz  # PyMuPDF：扫描型 PDF 渲染/裁剪
import pdfplumber
import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image

from ... import config
from .models import SUPPORTED_OUTPUT_FORMATS

logger = logging.getLogger(__name__)

# ---- 任务计数器（线程安全）----
_lock = threading.Lock()
_active_tasks = 0
_total_success_count = 0
_total_fail_count = 0
_total_elapsed_ms = 0


class DocConvertError(ValueError):
    """业务错误，携带建议的 HTTP 状态码。"""

    def __init__(self, detail: str, status: int = 400):
        super().__init__(detail)
        self.status = status
        self.detail = detail


def task_started() -> None:
    global _active_tasks
    with _lock:
        _active_tasks += 1


def task_finished() -> None:
    global _active_tasks
    with _lock:
        _active_tasks = max(0, _active_tasks - 1)


def record_conversion_metrics(success: bool, elapsed_ms: float) -> None:
    global _total_success_count, _total_fail_count, _total_elapsed_ms
    with _lock:
        if success:
            _total_success_count += 1
        else:
            _total_fail_count += 1
        _total_elapsed_ms += elapsed_ms


def get_metrics() -> Dict[str, Any]:
    with _lock:
        success = _total_success_count
        fail = _total_fail_count
        elapsed_tot = _total_elapsed_ms
        active = _active_tasks
    return {
        "active_tasks": active,
        "total_success_count": success,
        "total_fail_count": fail,
        "total_elapsed_ms": elapsed_tot,
        "average_elapsed_ms": round(elapsed_tot / success, 1) if success > 0 else 0.0,
    }


# ---- 路由表 ----
LOCAL_ONLY_EXTENSIONS = {
    ".docx", ".xlsx", ".pptx", ".txt", ".csv",
    ".html", ".htm", ".xml", ".json", ".md",
    ".zip", ".epub", ".rtf", ".odt", ".ods", ".odp",
}
OCR_ONLY_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
PDF_EXTENSION = ".pdf"

# ---- MarkItDown 实例（懒加载）----
_markitdown_instance = None
_markitdown_lock = threading.Lock()


def _get_markitdown():
    global _markitdown_instance
    if _markitdown_instance is not None:
        return _markitdown_instance
    with _markitdown_lock:
        if _markitdown_instance is None:
            from markitdown import MarkItDown
            _markitdown_instance = MarkItDown()
    return _markitdown_instance


def _markitdown_convert(stream: io.BytesIO, ext: str):
    """兼容不同版本的 MarkItDown API（convert_stream / convert）。"""
    md = _get_markitdown()
    try:
        return md.convert_stream(stream, file_extension=ext)
    except (TypeError, AttributeError):
        return md.convert(stream=stream, file_extension=ext)


def md_to_text(result) -> str:
    """从 MarkItDown 结果里取文本内容。"""
    text = getattr(result, "text_content", "")
    if text is None:
        return ""
    return text if isinstance(text, str) else str(text)


# ---- 本地 PaddleOCR 懒加载（可选依赖，用于 OCR 节点不可达时的降级）----
_local_ocr_instance = None
_local_ocr_init_lock = threading.Lock()
_local_ocr_exec_lock = threading.Lock()  # PaddleOCR 推理串行化，防止多线程抢占 C++ 状态
_local_ocr_init_error = None


def get_local_ocr():
    """获取本地 PaddleOCR 单例（CPU 模式）。未安装 paddleocr 时返回 None。"""
    global _local_ocr_instance, _local_ocr_init_error
    if _local_ocr_instance is not None:
        return _local_ocr_instance
    if _local_ocr_init_error is not None:
        return None
    with _local_ocr_init_lock:
        if _local_ocr_instance is not None:
            return _local_ocr_instance
        if _local_ocr_init_error is not None:
            return None
        try:
            from paddleocr import PaddleOCR
            logger.info("正在初始化本地 PaddleOCR（首次加载模型，CPU 模式）...")
            _local_ocr_instance = PaddleOCR(lang="ch", use_angle_cls=False, use_gpu=False)
            logger.info("本地 PaddleOCR 初始化完成（CPU 模式）")
            return _local_ocr_instance
        except Exception as exc:
            _local_ocr_init_error = str(exc)
            logger.error("本地 PaddleOCR 初始化失败: %s", exc)
            return None


# ---- PDF 分类 ----
def classify_pdf(file_bytes: bytes) -> str:
    """判断 PDF 是文本型还是扫描型：提取前 3 页文字，字符数 ≥ 阈值 → text。"""
    try:
        total_chars = 0
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i in range(min(3, len(pdf.pages))):
                page_text = pdf.pages[i].extract_text()
                if page_text:
                    total_chars += len(re.sub(r"\s+", "", page_text))
        return "text" if total_chars >= config.PDF_TEXT_THRESHOLD else "scanned"
    except Exception:
        return "scanned"


# ---- OCR 节点转发 ----
def check_ocr_node() -> Optional[Dict[str, Any]]:
    """探活 OCR 节点，返回健康信息或 None。"""
    try:
        resp = requests.get(f"{config.OCR_NODE_URL}/health", timeout=2)
        return resp.json() if resp.ok else None
    except Exception:
        return None


def forward_to_ocr_node(file_bytes: bytes, filename: str,
                        include_lines: bool = True) -> Tuple[Dict[str, Any], int]:
    """转发文件到 OCR 节点 /convert。返回 (response_dict, http_status)。"""
    try:
        resp = requests.post(
            f"{config.OCR_NODE_URL}/convert",
            files={"file": (filename, io.BytesIO(file_bytes))},
            params={"include_lines": "true" if include_lines else "false"},
            timeout=config.OCR_TIMEOUT,
        )
        try:
            return resp.json(), resp.status_code
        except ValueError:
            return {
                "success": False,
                "error": f"OCR 节点返回了非 JSON 响应（HTTP {resp.status_code}）",
                "filename": filename,
            }, 502
    except requests.exceptions.ConnectionError:
        return {
            "success": False,
            "error": f"OCR 节点不可达: {config.OCR_NODE_URL}",
            "filename": filename,
        }, 503
    except requests.exceptions.Timeout:
        return {
            "success": False,
            "error": f"OCR 节点处理超时（{config.OCR_TIMEOUT}s）",
            "filename": filename,
        }, 504


# ---- 裁剪 ----
CROP_EPS = 1e-4


def parse_crop(raw: Union[str, None]) -> Optional[Dict[str, float]]:
    """解析并校验 crop 表单字段。合法返回 {x,y,width,height} 比例值，缺省返回 None。"""
    if not raw:
        return None
    try:
        data = json.loads(raw)
    except (TypeError, ValueError):
        raise DocConvertError("crop 不是合法的 JSON")

    if not isinstance(data, dict):
        raise DocConvertError("crop 必须是 JSON 对象")
    if data.get("unit") != "ratio":
        raise DocConvertError('crop.unit 必须是 "ratio"')

    out: Dict[str, float] = {}
    for key in ("x", "y", "width", "height"):
        val = data.get(key)
        if not isinstance(val, (int, float)) or isinstance(val, bool):
            raise DocConvertError(f"crop.{key} 必须是数字")
        if val < 0 or val > 1:
            raise DocConvertError(f"crop.{key} 必须在 0~1 之间")
        out[key] = float(val)

    if out["width"] <= 0 or out["height"] <= 0:
        raise DocConvertError("crop.width / crop.height 必须大于 0")
    if out["x"] + out["width"] > 1 + CROP_EPS:
        raise DocConvertError("crop.x + crop.width 超过 1.0")
    if out["y"] + out["height"] > 1 + CROP_EPS:
        raise DocConvertError("crop.y + crop.height 超过 1.0")

    return out


def crop_image_bytes(file_bytes: bytes, crop: Dict[str, float]) -> bytes:
    """按比例裁剪图片，返回裁剪后的图片字节（保留原格式）。"""
    img = Image.open(io.BytesIO(file_bytes))
    w, h = img.size
    box = (
        round(crop["x"] * w),
        round(crop["y"] * h),
        round((crop["x"] + crop["width"]) * w),
        round((crop["y"] + crop["height"]) * h),
    )
    box = (
        min(box[0], w - 1),
        min(box[1], h - 1),
        max(box[2], box[0] + 1),
        max(box[3], box[1] + 1),
    )
    cropped = img.crop(box)
    buf = io.BytesIO()
    cropped.save(buf, format=(img.format or "PNG"))
    return buf.getvalue()


def crop_pdf_text(file_bytes: bytes, crop: Dict[str, float]) -> str:
    """文本型 PDF：按比例裁剪每一页并直接提取文字层，不走 OCR。"""
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            w, h = page.width, page.height
            bbox = (
                crop["x"] * w,
                crop["y"] * h,
                (crop["x"] + crop["width"]) * w,
                (crop["y"] + crop["height"]) * h,
            )
            try:
                text = page.crop(bbox).extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts)


def crop_pdf_page_images(file_bytes: bytes, crop: Dict[str, float]) -> List[Tuple[int, bytes]]:
    """扫描型 PDF：按比例裁剪每一页并渲染成图片。返回 [(page_no, png_bytes), ...]。"""
    images: List[Tuple[int, bytes]] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        mat = fitz.Matrix(config.CROP_PDF_ZOOM, config.CROP_PDF_ZOOM)
        for i, page in enumerate(doc, start=1):
            rect = page.rect
            clip = fitz.Rect(
                crop["x"] * rect.width,
                crop["y"] * rect.height,
                (crop["x"] + crop["width"]) * rect.width,
                (crop["y"] + crop["height"]) * rect.height,
            )
            pix = page.get_pixmap(matrix=mat, clip=clip)
            images.append((i, pix.tobytes("png")))
    finally:
        doc.close()
    return images


def _extract_ocr_text(result: Dict[str, Any]) -> str:
    """从 OCR 节点响应里尽量取出文字内容。"""
    for key in ("content", "text", "markdown"):
        val = result.get(key)
        if isinstance(val, str) and val.strip():
            return val.strip()
    return ""


# ---- 版面重建：基于行级 bbox 的几何启发式段落/标题分组 ----
# 只用三个廉价信号：行高（字号代理）→ 标题级别；行间距 → 段落分界；
# 左边界缩进 → 引用块。不做真正的版面分析（多栏/表格），阈值见 config。

_ASCII_WORD_RE = re.compile(r"[A-Za-z0-9]")


def _smart_join(prev_text: str, next_text: str) -> str:
    """同一段内跨行拼接。前后都是 ASCII 字母/数字时补空格，避免英文单词粘连。"""
    if not prev_text:
        return next_text
    if not next_text:
        return prev_text
    if _ASCII_WORD_RE.match(prev_text[-1]) and _ASCII_WORD_RE.match(next_text[0]):
        return prev_text + " " + next_text
    return prev_text + next_text


def _line_metrics(line: Dict[str, Any]) -> Dict[str, float]:
    """从一行 OCR 结果的 bbox 提取几何指标（top/bottom/height/left）。"""
    xs = [pt[0] for pt in line["bbox"]]
    ys = [pt[1] for pt in line["bbox"]]
    top, bottom = min(ys), max(ys)
    return {
        "text": (line.get("text") or "").strip(),
        "top": top,
        "bottom": bottom,
        "height": max(bottom - top, 1e-6),
        "left": min(xs),
    }


def _reconstruct_page_layout(page_lines: List[Dict[str, Any]]) -> str:
    """单页行列表 → 分组（标题/段落）→ 渲染成 Markdown 文本。"""
    metrics = [_line_metrics(ln) for ln in page_lines if ln.get("text", "").strip()]
    if not metrics:
        return ""

    heights = sorted(m["height"] for m in metrics)
    median_height = heights[len(heights) // 2] or 1.0

    body_lefts = [
        m["left"] for m in metrics
        if abs(m["height"] - median_height) <= 0.2 * median_height
    ]
    baseline_left = min(body_lefts) if body_lefts else min(m["left"] for m in metrics)

    groups: List[List[Dict[str, float]]] = []
    current: List[Dict[str, float]] = []
    for m in metrics:
        is_heading_line = m["height"] >= median_height * config.LAYOUT_HEADING_HEIGHT_RATIO
        if is_heading_line:
            if current:
                groups.append(current)
                current = []
            groups.append([m])
            continue
        if not current:
            current.append(m)
            continue
        gap = m["top"] - current[-1]["bottom"]
        indent_delta = abs(m["left"] - current[0]["left"])
        if (gap > median_height * config.LAYOUT_PARAGRAPH_GAP_RATIO
                or indent_delta > median_height * config.LAYOUT_INDENT_RATIO):
            groups.append(current)
            current = [m]
        else:
            current.append(m)
    if current:
        groups.append(current)

    rendered: List[str] = []
    for group in groups:
        if len(group) == 1 and group[0]["height"] >= median_height * config.LAYOUT_HEADING_HEIGHT_RATIO:
            ratio = group[0]["height"] / median_height
            level = 1 if ratio >= 2.0 else (2 if ratio >= 1.6 else 3)
            rendered.append(f"{'#' * level} {group[0]['text']}")
            continue
        text = group[0]["text"]
        for m in group[1:]:
            text = _smart_join(text, m["text"])
        indent = group[0]["left"] - baseline_left
        if indent > median_height * config.LAYOUT_INDENT_RATIO:
            text = f"> {text}"
        rendered.append(text)

    return "\n\n".join(rendered)


def _reconstruct_layout(lines: List[Dict[str, Any]],
                        force_page_markers: bool = False) -> str:
    """多页行列表 → 按页重建，页间用 "<!-- page N -->" 分隔。"""
    if not lines:
        return ""
    pages: Dict[int, List[Dict[str, Any]]] = {}
    for ln in lines:
        pages.setdefault(ln.get("page", 1), []).append(ln)

    use_markers = force_page_markers or len(pages) > 1
    parts: List[str] = []
    for page_no in sorted(pages.keys()):
        page_md = _reconstruct_page_layout(pages[page_no])
        if use_markers:
            parts.append(f"<!-- page {page_no} -->" + (f"\n{page_md}" if page_md else ""))
        elif page_md:
            parts.append(page_md)
    return "\n\n".join(p for p in parts if p)


def _apply_layout_reconstruction(result: Dict[str, Any],
                                 force_page_markers: bool = False) -> None:
    """就地改写 OCR 响应：按 bbox 重建 content，原始文本保留到 content_raw。"""
    result["layout_reconstructed"] = False
    if not config.LAYOUT_RECONSTRUCTION_ENABLED:
        result.pop("lines", None)
        return
    lines = result.get("lines")
    if not isinstance(lines, list) or not lines:
        result.pop("lines", None)
        return
    try:
        reconstructed = _reconstruct_layout(lines, force_page_markers=force_page_markers)
    except Exception as exc:
        result["layout_reconstruction_error"] = str(exc)
        result.pop("lines", None)
        return
    if reconstructed.strip():
        result["content_raw"] = result.get("content", "")
        result["content"] = reconstructed
        result["layout_reconstructed"] = True
    result.pop("lines", None)


def _extract_lines(ocr_result, page_num: Optional[int] = None) -> List[Dict[str, Any]]:
    """从 PaddleOCR 原始输出提取结构化行信息（文字 + bbox + 置信度）。"""
    if not ocr_result or len(ocr_result) == 0 or ocr_result[0] is None:
        return []
    lines = []
    for item in ocr_result[0]:
        if item is None:
            continue
        try:
            bbox_raw, (text, confidence) = item
            if not text or not str(text).strip():
                continue
            bbox = [[round(float(x), 1), round(float(y), 1)] for x, y in bbox_raw]
            line = {
                "text": str(text),
                "bbox": bbox,
                "confidence": round(float(confidence), 4),
            }
            if page_num is not None:
                line["page"] = page_num
            lines.append(line)
        except (IndexError, TypeError, ValueError):
            continue
    return lines


def _lines_to_text(lines: List[Dict[str, Any]]) -> str:
    return "\n".join(line["text"] for line in lines)


def _ocr_image_bytes(file_bytes: bytes, ocr, page_num: Optional[int] = None) -> List[Dict[str, Any]]:
    """对图片字节做本地 PaddleOCR 识别，返回行列表。"""
    import numpy as np
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode != "RGB":
        img = img.convert("RGB")
    img_array = np.array(img)
    with _local_ocr_exec_lock:
        raw_result = ocr.ocr(img_array, cls=False)
    return _extract_lines(raw_result, page_num=page_num)


def _do_ocr_local(file_bytes: bytes, ext: str, file_type: str) -> Dict[str, Any]:
    """本地 PaddleOCR CPU 模式全内存转换。"""
    ocr = get_local_ocr()
    if ocr is None:
        raise RuntimeError(f"本地 OCR 引擎不可用: {_local_ocr_init_error or '未知原因'}")

    if file_type == "image":
        page_lines = _ocr_image_bytes(file_bytes, ocr)
        return {"text": _lines_to_text(page_lines), "pages": 1, "lines": page_lines}

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = doc.page_count
    all_texts: List[str] = []
    all_lines: List[Dict[str, Any]] = []
    try:
        for page_num in range(page_count):
            pix = doc[page_num].get_pixmap(dpi=150)
            page_lines = _ocr_image_bytes(pix.tobytes("png"), ocr, page_num=page_num + 1)
            page_text = _lines_to_text(page_lines)
            if page_text.strip():
                all_texts.append(f"--- 第 {page_num + 1} 页 ---\n{page_text}")
            all_lines.extend(page_lines)
    finally:
        doc.close()
    return {"text": "\n\n".join(all_texts), "pages": page_count, "lines": all_lines}


def _convert_ocr_cropped_pdf(file_bytes: bytes, filename: str,
                             crop: Dict[str, float], t_start: float) -> Dict[str, Any]:
    """扫描型 PDF + 裁剪：逐页裁剪渲染后转 OCR 节点，拼接结果；支持本地 OCR 降级。"""
    ocr_node_up = check_ocr_node() is not None
    try:
        page_images = crop_pdf_page_images(file_bytes, crop)
    except Exception as exc:
        raise DocConvertError(f"渲染裁剪后的 PDF 页面失败: {exc}", 500)

    page_texts: List[str] = []
    page_errors: List[Dict[str, Any]] = []
    all_lines: List[Dict[str, Any]] = []

    if not ocr_node_up:
        logger.info("OCR 节点不可达，降级为本地 OCR 识别裁剪后的扫描 PDF: %s", filename)
        ocr = get_local_ocr()
        if ocr is None:
            raise DocConvertError(
                f"OCR 节点不可达且本地 OCR 不可用: {_local_ocr_init_error or 'paddleocr 未安装'}", 503
            )
        for page_no, img_bytes in page_images:
            try:
                page_lines = _ocr_image_bytes(img_bytes, ocr, page_num=page_no)
                page_texts.append(f"<!-- page {page_no} -->\n{_lines_to_text(page_lines)}")
                all_lines.extend(page_lines)
            except Exception as exc:
                page_errors.append({"page": page_no, "error": f"本地 OCR 失败: {exc}"})
    else:
        for page_no, img_bytes in page_images:
            result, status = forward_to_ocr_node(
                img_bytes, f"{filename}_page{page_no}.png")
            if result.get("success"):
                page_texts.append(f"<!-- page {page_no} -->\n{_extract_ocr_text(result)}")
                for ln in result.get("lines") or []:
                    ln["page"] = page_no
                    all_lines.append(ln)
            else:
                page_errors.append({"page": page_no, "error": result.get("error", f"HTTP {status}")})

    if not page_texts:
        raise DocConvertError("所有裁剪页 OCR 均失败", 502)

    resp = {
        "success": True,
        "filename": filename,
        "content": "\n\n".join(page_texts),
        "engine": "paddleocr_cpu_local" if not ocr_node_up else "paddleocr_cpu",
        "routed_to": "local_ocr" if not ocr_node_up else "ocr_node",
        "crop_applied": True,
        "pages_processed": len(page_images),
        "lines": all_lines,
    }
    _apply_layout_reconstruction(resp, force_page_markers=True)
    if page_errors:
        resp["page_errors"] = page_errors
    return resp


# ---- 输出格式：md / txt / docx ----
_INLINE_TOKEN_RE = re.compile(
    r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|__.+?__|_.+?_|`[^`]+?`|'
    r'!\[[^\]]*\]\([^)]*\)|\[[^\]]+\]\([^)]*\))'
)
_TABLE_SEP_CELL_RE = re.compile(r"^:?-{2,}:?$")


def validate_output_format(raw: str) -> str:
    fmt = (raw or "md").strip().lower()
    if fmt not in SUPPORTED_OUTPUT_FORMATS:
        raise DocConvertError(f"output_format 必须是 {sorted(SUPPORTED_OUTPUT_FORMATS)} 之一")
    return fmt


def strip_markdown_light(md_text: str) -> str:
    """剥离常见 Markdown 语法记号，得到接近纯文本的内容（用于 txt 输出）。"""
    text = md_text or ""
    text = re.sub(r"^(#{1,6})\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^```[^\n]*\n", "", text, flags=re.MULTILINE)
    text = re.sub(r"^```\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"(\*\*\*|\*\*|\*|__|_)(.+?)\1", r"\2", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    plain_lines = []
    for _line in text.splitlines():
        _stripped = _line.strip()
        if _stripped.startswith("|"):
            _cells = [c.strip() for c in _stripped.strip("|").split("|")]
            if all(_TABLE_SEP_CELL_RE.fullmatch(c) for c in _cells if c):
                continue
            if not any(_cells):
                continue
            plain_lines.append("   ".join(_cells))
        else:
            plain_lines.append(_line)
    text = "\n".join(plain_lines)
    text = re.sub(r"^(\s*[-*_]\s*){3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _add_inline_runs(paragraph, text: str) -> None:
    """把一段文本里的 **加粗**、*斜体*、`行内代码`、![图](url)、[链接](url) 写成 docx run。"""
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("!["):
            alt = re.match(r"!\[([^\]]*)\]", token).group(1)
            if alt:
                r = paragraph.add_run(alt)
                r.italic = True
        elif token.startswith("[") and token.endswith(")"):
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            r = paragraph.add_run(label)
            r.underline = True
        elif token.startswith("***") and token.endswith("***"):
            r = paragraph.add_run(token[3:-3])
            r.bold = True
            r.italic = True
        elif token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
        else:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _split_table_row(line: str) -> Optional[List[str]]:
    s = line.strip()
    if not s.startswith("|"):
        return None
    return [c.strip() for c in s.strip("|").split("|")]


def _set_docx_default_font(doc) -> None:
    """把 docx 默认样式设为常见中文字体，避免缺字体导致的渲染问题。"""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_docx_hr(doc) -> None:
    """在文档里插入一条水平分割线（底部边框实现）。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx_bytes_fallback(markdown_text: str) -> bytes:
    """用 python-docx 手写 Markdown → docx（pandoc 缺失时的兜底引擎）。
    支持：标题、表格、无序/有序列表、引用块、分割线、代码块、行内样式。"""
    doc = Document()
    _set_docx_default_font(doc)
    lines = (markdown_text or "").splitlines()

    in_code_block = False
    code_lines: List[str] = []

    def flush_code_block():
        if not code_lines:
            return
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_lines))
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Pt(18)
        code_lines.clear()

    idx = 0
    n = len(lines)
    while idx < n:
        line = lines[idx].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue
        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            doc.add_paragraph("")
            idx += 1
            continue

        h_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if h_match:
            doc.add_heading(h_match.group(2).strip(), level=min(len(h_match.group(1)), 9))
            idx += 1
            continue

        if _split_table_row(line) is not None:
            table_lines = [line]
            idx += 1
            while idx < n and _split_table_row(lines[idx].rstrip("\n")) is not None:
                table_lines.append(lines[idx].rstrip("\n"))
                idx += 1
            rows: List[List[str]] = []
            for tl in table_lines:
                cells = _split_table_row(tl) or []
                if all(_TABLE_SEP_CELL_RE.fullmatch(c) for c in cells if c):
                    continue
                rows.append(cells)
            if rows:
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        _add_inline_runs(tbl.cell(ri, ci).paragraphs[0], cell_text)
                doc.add_paragraph("")
            continue

        bq_match = re.match(r"^(>+)\s?(.*)$", stripped)
        if bq_match:
            p = doc.add_paragraph()
            _add_inline_runs(p, bq_match.group(2))
            p.paragraph_format.left_indent = Pt(18 + 12 * (len(bq_match.group(1)) - 1))
            p.paragraph_format.space_after = Pt(2)
            idx += 1
            continue

        if re.fullmatch(r"(\s*[-*_]\s*){3,}\s*", stripped):
            _add_docx_hr(doc)
            idx += 1
            continue

        ul_match = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, ul_match.group(1))
            idx += 1
            continue

        ol_match = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ol_match.group(1))
            idx += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        idx += 1

    if in_code_block:
        flush_code_block()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _markdown_to_docx_pandoc(markdown_text: str) -> bytes:
    """用 pandoc 做 Markdown → docx（更完整的语义与样式，可套用参考模板）。"""
    pandoc = shutil.which("pandoc")
    cmd = [pandoc, "-f", "markdown", "-t", "docx"]
    if config.DOCX_REFERENCE_TEMPLATE and os.path.isfile(config.DOCX_REFERENCE_TEMPLATE):
        cmd.append(f"--reference-doc={config.DOCX_REFERENCE_TEMPLATE}")
    cmd += ["-o", "-"]
    proc = subprocess.run(
        cmd,
        input=(markdown_text or "").encode("utf-8"),
        capture_output=True,
        timeout=config.PANDOC_TIMEOUT,
    )
    if proc.returncode != 0:
        stderr = proc.stderr.decode("utf-8", errors="ignore").strip()
        raise RuntimeError(f"pandoc 退出码 {proc.returncode}: {stderr[:500]}")
    if not proc.stdout:
        raise RuntimeError("pandoc 输出为空")
    return proc.stdout


def markdown_to_docx_bytes(markdown_text: str) -> Tuple[bytes, str]:
    """Markdown → docx 统一入口。返回 (docx_bytes, engine_name)。
    pandoc 优先；未安装或转换失败时降级为 python-docx。"""
    pandoc = shutil.which("pandoc")
    if pandoc:
        try:
            return _markdown_to_docx_pandoc(markdown_text), "pandoc"
        except Exception as exc:
            logger.warning("pandoc 转换失败，降级为 python-docx: %s", exc)
    return markdown_to_docx_bytes_fallback(markdown_text), "python-docx-fallback"


# ---- 本地转换（MarkItDown）----
def convert_local(file_bytes: bytes, filename: str, ext: str,
                  t_start: float) -> Dict[str, Any]:
    result = _markitdown_convert(io.BytesIO(file_bytes), ext)
    return {
        "success": True,
        "filename": filename,
        "file_size_mb": round(len(file_bytes) / (1024 * 1024), 2),
        "content": md_to_text(result),
        "engine": "markitdown",
        "routed_to": "local",
        "elapsed_ms": round((time.perf_counter() - t_start) * 1000),
    }


# ---- OCR 转换 ----
def convert_ocr(file_bytes: bytes, filename: str, t_start: float) -> Dict[str, Any]:
    """图片/扫描 PDF → OCR 节点；节点不可达时降级本地 PaddleOCR。"""
    ocr_node_up = check_ocr_node() is not None
    if not ocr_node_up:
        logger.info("OCR 节点不可达，降级为本地 OCR 识别: %s", filename)
        ext = Path(filename).suffix.lower()
        file_type = "pdf" if ext == PDF_EXTENSION else "image"
        try:
            result = _do_ocr_local(file_bytes, ext, file_type)
            resp = {
                "success": True,
                "filename": filename,
                "content": result["text"],
                "engine": "paddleocr_cpu_local",
                "routed_to": "local_ocr",
                "pages": result["pages"],
                "lines": result["lines"],
                "elapsed_ms": round((time.perf_counter() - t_start) * 1000),
            }
            _apply_layout_reconstruction(resp)
            return resp
        except Exception as exc:
            logger.error("本地 OCR 降级转换失败: %s", exc)
            if ext == PDF_EXTENSION:
                try:
                    return convert_local(file_bytes, filename, ext, t_start)
                except Exception:
                    pass
            raise DocConvertError(f"OCR 节点不可达且本地 OCR 失败: {exc}", 503)

    result, status = forward_to_ocr_node(file_bytes, filename)
    if result.get("success"):
        result["routed_to"] = "ocr_node"
        _apply_layout_reconstruction(result)
    result["route_elapsed_ms"] = round((time.perf_counter() - t_start) * 1000)
    if not result.get("success"):
        raise DocConvertError(result.get("error", "OCR 转换失败"),
                              status if 400 <= status <= 599 else 502)
    return result


# ---- 统一转换入口 ----
def convert_document(content: bytes, filename: str, crop: Optional[Dict[str, float]],
                     output_format: str = "md", t_start: Optional[float] = None) -> Dict[str, Any]:
    """统一转换入口。按扩展名自动路由，返回转换结果 dict（含 content / base64 内容）。"""
    output_format = validate_output_format(output_format)
    t_start = t_start or time.perf_counter()
    task_started()
    try:
        file_size_mb = len(content) / (1024 * 1024)
        if file_size_mb > config.CONVERT_MAX_MB:
            raise DocConvertError(
                f"文件过大（{file_size_mb:.1f}MB），上限 {config.CONVERT_MAX_MB}MB", 413)

        ext = Path(filename).suffix.lower()
        if not ext:
            raise DocConvertError("文件缺少扩展名，无法判断格式")
        if ext == ".docx" and output_format == "docx":
            raise DocConvertError(
                "源文件是 .docx，再选 docx 输出没有意义（会先转 Markdown 再重建，"
                "原始排版/图片/格式必然丢失）；请直接使用原文件，或改选 md/txt")

        if ext in LOCAL_ONLY_EXTENSIONS:
            result = convert_local(content, filename, ext, t_start)
            if crop is not None:
                result["crop_ignored"] = "办公格式不支持区域裁剪，已转换整份文件"
        elif ext in OCR_ONLY_EXTENSIONS:
            if crop is not None:
                try:
                    content = crop_image_bytes(content, crop)
                except Exception as exc:
                    raise DocConvertError(f"裁剪图片失败: {exc}")
            result = convert_ocr(content, filename, t_start)
            if crop is not None and result.get("success"):
                result["crop_applied"] = True
        elif ext == PDF_EXTENSION:
            if classify_pdf(content) == "text":
                if crop is not None:
                    result = {
                        "success": True,
                        "filename": filename,
                        "content": crop_pdf_text(content, crop),
                        "engine": "pdfplumber_crop",
                        "routed_to": "local",
                        "crop_applied": True,
                    }
                else:
                    result = convert_local(content, filename, ext, t_start)
            else:
                if crop is not None:
                    result = _convert_ocr_cropped_pdf(content, filename, crop, t_start)
                else:
                    result = convert_ocr(content, filename, t_start)
        else:
            raise DocConvertError(f"不支持的文件格式 '{ext}'")

        result = apply_output_format(result, output_format, filename)
        record_conversion_metrics(True, round((time.perf_counter() - t_start) * 1000))
        return result
    except DocConvertError:
        record_conversion_metrics(False, round((time.perf_counter() - t_start) * 1000))
        raise
    finally:
        task_finished()


def apply_output_format(result: Dict[str, Any], output_format: str,
                        filename: str) -> Dict[str, Any]:
    """按 output_format 改写成功响应：md 原样、txt 剥离语法、docx 生成文件 base64。"""
    base_name = os.path.splitext(os.path.basename(filename or "converted"))[0] or "converted"
    content = result.get("content", "")
    result["output_format"] = output_format

    if output_format == "md":
        result["output_filename"] = f"{base_name}.md"
        result["mime_type"] = "text/markdown"
    elif output_format == "txt":
        result["output_filename"] = f"{base_name}.txt"
        result["mime_type"] = "text/plain"
        result["content"] = strip_markdown_light(content)
    elif output_format == "docx":
        try:
            docx_bytes, docx_engine = markdown_to_docx_bytes(content)
        except Exception as exc:
            raise DocConvertError(f"生成 docx 失败: {exc}", 500)
        result["output_filename"] = f"{base_name}.docx"
        result["mime_type"] = ("application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")
        result["file_base64"] = base64.b64encode(docx_bytes).decode("ascii")
        result["docx_engine"] = docx_engine
        result.pop("content", None)

    return result